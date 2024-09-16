import pdfplumber
from docx import Document
from dotenv import load_dotenv

from src.text_processing import convert_roman_to_arabic, expand_ranges_in_text
from src.load_env import load_environment_variables

def pdf_to_word_and_extract_table():
    # Load environment variables from .env file
    PDF_FILE_PATH, OUTPUT_WORD_PATH, OUTPUT_TXT_PATH, _ = load_environment_variables()

    # Define headers for each round of competition corresponding to each page
    headers = {
        1: "ОБЛАСТЕН КРЪГ",  # Page 2 in the PDF (index 1)
        2: "РЕГИОНАЛЕН КРЪГ",  # Page 3 in the PDF (index 2)
        3: "НАЦИОНАЛЕН КРЪГ"   # Page 4 in the PDF (index 3)
    }

    with pdfplumber.open(PDF_FILE_PATH) as pdf:
        doc = Document()
        doc.add_heading('Extracted Tables', 0)

        with open(OUTPUT_TXT_PATH, 'w', encoding='utf-8') as txt_file:
            # Loop through pages 2, 3, and 4 (indices 1, 2, 3)
            for page_index, header in headers.items():
                page = pdf.pages[page_index]  # Fetch the current page
                tables = page.extract_tables()

                if tables:
                    target_table = tables[0]

                    # Add header to the Word document and text file
                    doc.add_paragraph(header)  # For Word document
                    txt_file.write(header + '\n')  # For TXT file

                    # Add the table to the Word document
                    table_in_doc = doc.add_table(rows=1, cols=len(target_table[0]))
                    hdr_cells = table_in_doc.rows[0].cells
                    for i, header in enumerate(target_table[0]):
                        hdr_cells[i].text = convert_roman_to_arabic(header)

                    for row in target_table[1:]:
                        row_cells = table_in_doc.add_row().cells
                        buffer = ""
                        for i, cell in enumerate(row):
                            formatted_cell = convert_roman_to_arabic(cell).strip().replace('\n', ' ')
                            expanded_cell = expand_ranges_in_text(formatted_cell)
                            row_cells[i].text = formatted_cell  # Add to Word doc
                            buffer += expanded_cell + "\t"

                            # Write to the TXT file after a complete row
                            if "Г." in expanded_cell:
                                txt_file.write(buffer.strip() + '\n')
                                buffer = ""

                        # In case the buffer wasn't written due to missing "Г."
                        if buffer.strip():
                            txt_file.write(buffer.strip() + '\n')

            # Save the Word document
            doc.save(OUTPUT_WORD_PATH)
            print(f"Word document saved successfully: {OUTPUT_WORD_PATH}")
            print(f"Text file saved successfully: {OUTPUT_TXT_PATH}")


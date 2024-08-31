import pdfplumber
from docx import Document
import re
import json
import os
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Retrieve file paths from environment variables
PDF_FILE_PATH = os.getenv('PDF_FILE_PATH')
OUTPUT_WORD_PATH = os.getenv('OUTPUT_WORD_PATH')
OUTPUT_TXT_PATH = os.getenv('OUTPUT_TXT_PATH')
OUTPUT_JSON_PATH = os.getenv('OUTPUT_JSON_PATH')

# Expanded dictionary for Roman to Arabic numeral conversion
roman_to_arabic = {
    "I": 1, "II": 2, "III": 3, "IV": 4, "V": 5, "VI": 6, "VII": 7, "VIII": 8,
    "IX": 9, "X": 10, "XI": 11, "XII": 12
}

# Translation dictionary for the subjects
translation_dict = {
    "ЗНАМ И МОГА": "know_and_can",
    "МАТЕМАТИКА": "math",
    "БЪЛГАРСКИ ЕЗИК И ЛИТЕРАТУРА": "bulgarian_language_and_literature",
    "АНГЛИЙСКИ ЕЗИК": "english",
    "НЕМСКИ ЕЗИК": "german",
    "ИСПАНСКИ ЕЗИК": "spanish",
    "ИТАЛИАНСКИ ЕЗИК": "italian",
    "РУСКИ ЕЗИК": "russian",
    "ФРЕНСКИ ЕЗИК": "french",
    "ИНФОРМАТИКА": "informatics",
    "ИНФОРМАЦИОННИ ТЕXНОЛОГИИ": "information_technologies",
    "ЛИНГВИСТИКА": "linguistics",
    "ФИЛОСОФИЯ": "philosophy",
    "ИСТОРИЯ И ЦИВИЛИЗАЦИИ": "history_and_civilizations",
    "ГЕОГРАФИЯ И ИКОНОМИКА": "geography_and_economics",
    "ГРАЖДАНСКО ОБРАЗОВАНИЕ": "civic_education",
    "ФИЗИКА": "physics",
    "АСТРОНОМИЯ": "astronomy",
    "XИМИЯ И ОПАЗВАНЕ НА ОКОЛНАТА СРЕДА": "chemistry_and_environmental_protection",
    "БИОЛОГИЯ И ЗДРАВНО ОБРАЗОВАНИЕ": "biology_and_health_education",
    "ТЕXНИЧЕСКО ЧЕРТАНЕ": "technical_drawing"
}


def expand_range(start, end):
    return ' '.join(map(str, range(start, end + 1)))


# Function to process the entire text and expand all ranges
def expand_ranges_in_text(text):
    # Regex pattern to find ranges
    pattern = re.compile(r'(\d+)\s*-\s*(\d+)')

    # Find all ranges and expand them
    matches = pattern.finditer(text)
    expanded_text = text

    for match in matches:
        start, end = map(int, match.groups())
        expanded_text = expanded_text.replace(match.group(0), expand_range(start, end))

    return expanded_text


def normalize_text(text):
    # Normalize the text to handle Cyrillic letters and case differences
    text = text.upper()
    text = text.replace("Х", "X").replace("І", "I").replace("–", "-")
    return text


def convert_roman_to_arabic(text):
    text = normalize_text(text)
    # Replace Roman numerals in text with Arabic numbers
    for roman, arabic in roman_to_arabic.items():
        text = re.sub(rf'\b{roman}\b', str(arabic), text)
    return text


def clean_date(date_str):
    # Remove "ДО" and "Г." and keep only the date in format DD.MM.YYYY
    date_str = date_str.replace("ДО", "").replace("Г.", "").strip()
    return date_str


def extract_and_process_text():
    # Load and process the text file
    olympiad_data = {}

    with open(OUTPUT_TXT_PATH, 'r', encoding='utf-8') as file:
        for line in file:
            # Split the line into components (subject, class, date)
            parts = line.split("\t")
            if len(parts) < 4:
                continue

            subject = parts[0].strip()
            classes_str = parts[1].strip()
            date_str = clean_date(parts[-1].strip())

            # Translate the subject
            translated_subject = translation_dict.get(subject, subject)

            # Replace Roman numerals with Arabic numbers and process ranges
            classes_str = convert_roman_to_arabic(classes_str.replace("И", "").strip())
            classes = []

            class_parts = re.split(r'\s+', classes_str)  # Split by whitespace

            for part in class_parts:
                part = part.strip().replace(",", "")
                if "-" in part:
                    try:
                        start, end = map(int, part.split("-"))
                        classes.extend(range(start, end + 1))  # Create a range from start to end
                    except ValueError:
                        print(f"Warning: Unable to process class range '{part}'.")
                else:
                    try:
                        classes.append(int(part))
                    except ValueError:
                        print(f"Warning: Unable to process class part '{part}'.")

            # Remove duplicates and sort the list of classes
            classes = sorted(set(classes))

            # Add to the dictionary
            olympiad_data[translated_subject] = {
                "class": classes,
                "date": date_str
            }

    # Save to a JSON file
    with open(OUTPUT_JSON_PATH, 'w', encoding='utf-8') as json_file:
        json.dump(olympiad_data, json_file, ensure_ascii=False, indent=4)
    print(f"JSON file saved successfully: {OUTPUT_JSON_PATH}")


def pdf_to_word_and_extract_table():
    # Open the PDF file
    with pdfplumber.open(PDF_FILE_PATH) as pdf:
        # Extract text from page 2 (0-indexed)
        page = pdf.pages[1]  # Page 2 in 0-indexed format

        # Extract tables
        tables = page.extract_tables()

        # We assume the desired table is the first one
        if tables:
            target_table = tables[0]

            # Create a Word document
            doc = Document()
            doc.add_heading('Extracted Table', 0)

            # Add the table to the Word document
            table_in_doc = doc.add_table(rows=1, cols=len(target_table[0]))
            hdr_cells = table_in_doc.rows[0].cells

            # Add headers
            for i, header in enumerate(target_table[0]):
                hdr_cells[i].text = convert_roman_to_arabic(header)

            # Add rows to the Word document
            for row in target_table[1:]:
                row_cells = table_in_doc.add_row().cells
                for i, cell in enumerate(row):
                    row_cells[i].text = convert_roman_to_arabic(cell)

            # Save the Word document
            doc.save(OUTPUT_WORD_PATH)
            print(f"Word document saved successfully: {OUTPUT_WORD_PATH}")

            # Save the table to a text file with better formatting
            with open(OUTPUT_TXT_PATH, 'w', encoding='utf-8') as txt_file:
                buffer = ""
                for row in target_table[1:]:  # Skip the first row which is the header
                    for cell in row:
                        # Convert Roman to Arabic and expand ranges
                        formatted_cell = convert_roman_to_arabic(cell).strip().replace('\n', ' ')
                        # Removing the "-" symbol
                        expanded_cell = expand_ranges_in_text(formatted_cell)

                        buffer += expanded_cell + "\t"
                        # Only after "Г." there must be a new line
                        if "Г." in expanded_cell:
                            txt_file.write(buffer.strip() + '\n')
                            buffer = ""

                # Ensure the last line is written if not already
                if buffer.strip():
                    txt_file.write(buffer.strip() + '\n')

            print(f"Text file saved successfully: {OUTPUT_TXT_PATH}")

            # Process the text file and create the JSON
            extract_and_process_text()

        else:
            print("No tables found on page 2.")

